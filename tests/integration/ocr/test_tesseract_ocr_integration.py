from io import BytesIO
import os
from pathlib import Path
import time
import unittest

from docx import Document
from dotenv import load_dotenv
from PIL import Image, ImageDraw, ImageFont

from services.document_service import DocumentService


load_dotenv(".env")

RUN_OCR_INTEGRATION_TESTS = (
    os.getenv("RUN_OCR_INTEGRATION_TESTS", "").strip() == "1"
)


class UploadedRequirement(BytesIO):
    def __init__(self, name: str, payload: bytes):
        super().__init__(payload)
        self.name = name


@unittest.skipUnless(
    RUN_OCR_INTEGRATION_TESTS,
    "set RUN_OCR_INTEGRATION_TESTS=1 to run real Tesseract OCR tests",
)
class TesseractOcrIntegrationTests(unittest.TestCase):
    def test_real_tesseract_ocr_flows_through_document_service(self):
        executable = Path(os.environ["TESSERACT_CMD"])
        self.assertTrue(executable.is_file(), "Tesseract executable is missing")
        font_path = Path(
            os.getenv("OCR_TEST_FONT", r"C:\Windows\Fonts\msyh.ttc")
        )
        self.assertTrue(font_path.is_file(), "Chinese OCR test font is missing")

        expected = "商户单日提现上限为二十万元"
        image = Image.new("RGB", (1200, 180), "white")
        draw = ImageDraw.Draw(image)
        font = ImageFont.truetype(str(font_path), 44)
        draw.text((40, 55), expected, fill="black", font=font)
        image_buffer = BytesIO()
        image.save(image_buffer, format="PNG")
        image_buffer.seek(0)

        document = Document()
        document.add_paragraph().add_run().add_picture(image_buffer)
        document_buffer = BytesIO()
        document.save(document_buffer)

        started = time.perf_counter()
        content = DocumentService.parse(
            UploadedRequirement("ocr-smoke.docx", document_buffer.getvalue())
        )
        elapsed_seconds = time.perf_counter() - started

        normalized_text = "".join(content.to_plain_text().split())
        self.assertIn(expected, normalized_text)
        self.assertEqual(content.stats.image_count, 1)
        self.assertEqual(content.stats.ocr_element_count, 1)
        self.assertEqual(content.stats.low_confidence_ocr_count, 0)
        self.assertEqual(content.stats.failed_ocr_count, 0)
        self.assertFalse(content.warnings)
        self.assertGreater(elapsed_seconds, 0)


if __name__ == "__main__":
    unittest.main()
