import json
from pathlib import Path

from docx import Document
from PIL import Image
from pypdf import PdfReader


FIXTURE_DIR = Path(__file__).parents[3] / "evaluation" / "fixtures"
MANIFEST_PATH = FIXTURE_DIR / "gold_v1.json"


def _load_manifest() -> dict:
    with MANIFEST_PATH.open(encoding="utf-8") as file:
        return json.load(file)


def _fixtures_by_target() -> dict[str, dict]:
    return {
        item["evaluation_target"]: item
        for item in _load_manifest()["fixtures"]
    }


def test_document_fixture_manifest_has_five_unique_targets_and_valid_files():
    manifest = _load_manifest()
    fixtures = manifest["fixtures"]

    assert manifest["schema_version"] == 1
    assert len(fixtures) == 5
    assert len({item["fixture_id"] for item in fixtures}) == 5
    assert {item["evaluation_target"] for item in fixtures} == {
        "native_text",
        "table_structure",
        "ocr_text",
        "flow_semantics",
        "ui_semantics",
    }
    for item in fixtures:
        path = FIXTURE_DIR / item["path"]
        assert path.is_file()
        assert path.stat().st_size > 0


def test_native_pdf_contains_expected_extractable_text():
    fixture = _fixtures_by_target()["native_text"]
    text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(FIXTURE_DIR / fixture["path"]).pages
    )

    for expected_line in fixture["gold"]["text_lines"]:
        assert expected_line in text


def test_scanned_pdf_has_no_native_text_layer():
    fixture = _fixtures_by_target()["ocr_text"]
    text = "".join(
        page.extract_text() or ""
        for page in PdfReader(FIXTURE_DIR / fixture["path"]).pages
    )

    assert text.strip() == ""


def test_docx_contains_expected_text_and_table_structure():
    fixture = _fixtures_by_target()["table_structure"]
    document = Document(FIXTURE_DIR / fixture["path"])
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table = [[cell.text for cell in row.cells] for row in document.tables[0].rows]

    for expected_line in fixture["gold"]["text_lines"]:
        assert expected_line in paragraphs
    assert table[0] == fixture["gold"]["table"]["headers"]
    assert table[1:] == fixture["gold"]["table"]["rows"]


def test_image_fixtures_are_valid_and_have_expected_dimensions():
    fixtures = _fixtures_by_target()

    with Image.open(FIXTURE_DIR / fixtures["flow_semantics"]["path"]) as image:
        assert image.format == "PNG"
        assert image.size == (1200, 700)
    with Image.open(FIXTURE_DIR / fixtures["ui_semantics"]["path"]) as image:
        assert image.format == "PNG"
        assert image.size == (1200, 760)
