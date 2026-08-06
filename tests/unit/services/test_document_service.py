from io import BytesIO
import sys
from types import ModuleType, SimpleNamespace
import unittest
from unittest.mock import patch

from documents import (
    DocumentFormat,
    DocumentImageElement,
    DocumentOcrDisposition,
    DocumentOcrElement,
    DocumentVisualElement,
    DocumentVisualKind,
    DocumentVisualNode,
    DocumentVisualRelation,
    DocumentTableElement,
    DocumentTextElement,
    DocumentTextKind,
    DocumentWarningCode,
)
from services.document_service import DocumentService
from services.ocr_service import OcrError, OcrTextLine
from services.visual_service import VisualUnderstandingResult


class UploadedRequirement(BytesIO):
    def __init__(self, name: str, content: str | bytes):
        super().__init__(
            content if isinstance(content, bytes) else content.encode("utf-8")
        )
        self.name = name


class DocumentServiceTests(unittest.TestCase):
    def test_uploaded_text_requirement_is_extracted(self):
        uploaded = UploadedRequirement(
            "订单需求.txt",
            "用户提交订单时需要校验库存。",
        )

        result = DocumentService.extract_text(uploaded)

        self.assertEqual(result, "用户提交订单时需要校验库存。")

    def test_uploaded_markdown_requirement_is_extracted(self):
        uploaded = UploadedRequirement(
            "订单需求.md",
            "# 订单需求\n\n库存不足时禁止提交。",
        )

        result = DocumentService.extract_text(uploaded)

        self.assertEqual(
            result,
            "# 订单需求\n\n库存不足时禁止提交。",
        )

    def test_markdown_is_parsed_into_source_aware_ordered_elements(self):
        uploaded = UploadedRequirement(
            "订单需求.md",
            "# 订单需求\n\n- 校验库存\n- 扣减库存\n\n失败时提示用户。",
        )

        content = DocumentService.parse(uploaded)

        self.assertEqual(content.document_format, DocumentFormat.MARKDOWN)
        self.assertEqual(
            [element.kind for element in content.elements],
            [
                DocumentTextKind.TITLE,
                DocumentTextKind.LIST_ITEM,
                DocumentTextKind.LIST_ITEM,
                DocumentTextKind.PARAGRAPH,
            ],
        )
        self.assertEqual(
            [element.source.element_index for element in content.elements],
            [0, 1, 2, 3],
        )
        self.assertTrue(
            all(
                element.source.source_id.startswith(content.document_id)
                for element in content.elements
            )
        )

    def test_same_file_produces_stable_document_and_source_ids(self):
        first = DocumentService.parse(
            UploadedRequirement("需求.txt", "第一段\n\n第二段")
        )
        second = DocumentService.parse(
            UploadedRequirement("需求.txt", "第一段\n\n第二段")
        )

        self.assertEqual(first.document_id, second.document_id)
        self.assertEqual(
            [element.source.source_id for element in first.elements],
            [element.source.source_id for element in second.elements],
        )

    def test_empty_text_document_returns_explicit_warning(self):
        content = DocumentService.parse(UploadedRequirement("空需求.txt", ""))

        self.assertEqual(content.elements, ())
        self.assertEqual(
            [warning.code for warning in content.warnings],
            [DocumentWarningCode.EMPTY_DOCUMENT],
        )

    def test_pdf_keeps_page_numbers_and_reports_empty_page(self):
        class Pdf:
            pages = [
                SimpleNamespace(
                    extract_text=lambda: "第一页需求",
                    extract_tables=lambda: [],
                    curves=[],
                ),
                SimpleNamespace(
                    extract_text=lambda: "",
                    extract_tables=lambda: [],
                    curves=[],
                ),
                SimpleNamespace(
                    extract_text=lambda: "第三页规则",
                    extract_tables=lambda: [],
                    curves=[],
                ),
            ]

            def __enter__(self):
                return self

            def __exit__(self, *_):
                return None

        pdfplumber = ModuleType("pdfplumber")
        pdfplumber.open = lambda _: Pdf()
        pypdf = ModuleType("pypdf")
        pypdf.PdfReader = lambda _: SimpleNamespace(
            pages=[
                SimpleNamespace(images=[]),
                SimpleNamespace(images=[]),
                SimpleNamespace(images=[]),
            ]
        )

        with patch.dict(
            sys.modules, {"pdfplumber": pdfplumber, "pypdf": pypdf}
        ):
            content = DocumentService.parse(
                UploadedRequirement("需求.pdf", b"fake-pdf")
            )

        self.assertEqual(content.to_plain_text(), "第一页需求\n\n第三页规则")
        self.assertEqual(
            [element.source.page_number for element in content.elements],
            [1, 3],
        )
        self.assertEqual(content.warnings[0].code, DocumentWarningCode.EMPTY_PAGE)
        self.assertEqual(content.warnings[0].source.page_number, 2)

    def test_docx_extracts_ordered_paragraph_table_and_embedded_image(self):
        from docx import Document
        from docx.shared import Inches
        from PIL import Image

        image_buffer = BytesIO()
        Image.new("RGB", (8, 6), "blue").save(image_buffer, format="PNG")
        image_buffer.seek(0)
        document = Document()
        document.add_heading("支付需求", level=1)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "状态"
        table.cell(0, 1).text = "行为"
        table.cell(1, 0).text = "成功"
        table.cell(1, 1).text = "创建订单"
        paragraph = document.add_paragraph("支付流程图")
        paragraph.add_run().add_picture(image_buffer, width=Inches(1))
        payload = BytesIO()
        document.save(payload)

        content = DocumentService.parse(
            UploadedRequirement("需求.docx", payload.getvalue())
        )

        self.assertEqual(
            [type(element) for element in content.elements],
            [
                DocumentTextElement,
                DocumentTableElement,
                DocumentTextElement,
                DocumentImageElement,
            ],
        )
        self.assertEqual(content.elements[0].kind, DocumentTextKind.TITLE)
        self.assertEqual(content.elements[1].table.rows[1][1], "创建订单")
        self.assertEqual(content.elements[3].image.mime_type, "image/png")
        self.assertEqual(len(content.attachments), 1)
        self.assertIn("状态 | 行为", content.to_plain_text())
        self.assertEqual(content.stats.table_count, 1)
        self.assertEqual(content.stats.image_count, 1)

    def test_pdf_extracts_recognizable_table_image_and_vector_warning(self):
        class Pdf:
            pages = [
                SimpleNamespace(
                    extract_text=lambda: "退款规则",
                    extract_tables=lambda: [
                        [["状态", "行为"], ["成功", "原路退款"]]
                    ],
                    curves=[{"object_type": "curve"}],
                )
            ]

            def __enter__(self):
                return self

            def __exit__(self, *_):
                return None

        pdfplumber = ModuleType("pdfplumber")
        pdfplumber.open = lambda _: Pdf()
        pypdf = ModuleType("pypdf")
        pypdf.PdfReader = lambda _: SimpleNamespace(
            pages=[
                SimpleNamespace(
                    images=[SimpleNamespace(data=b"png", name="figure.png")]
                )
            ]
        )

        with patch.dict(
            sys.modules, {"pdfplumber": pdfplumber, "pypdf": pypdf}
        ):
            content = DocumentService.parse(
                UploadedRequirement("需求.pdf", b"fake-pdf")
            )

        self.assertEqual(content.stats.page_count, 1)
        self.assertEqual(content.stats.table_count, 1)
        self.assertEqual(content.stats.image_count, 1)
        self.assertEqual(content.elements[1].source.page_number, 1)
        self.assertEqual(content.elements[1].table.rows[1][1], "原路退款")
        self.assertEqual(content.attachments[0].content, b"png")
        self.assertIn(
            DocumentWarningCode.PAGE_RENDER_REQUIRED,
            [warning.code for warning in content.warnings],
        )

    def test_oversized_image_is_skipped_with_coverage_warning(self):
        class Pdf:
            pages = [
                SimpleNamespace(
                    extract_text=lambda: "图片需求",
                    extract_tables=lambda: [],
                    curves=[],
                )
            ]

            def __enter__(self):
                return self

            def __exit__(self, *_):
                return None

        pdfplumber = ModuleType("pdfplumber")
        pdfplumber.open = lambda _: Pdf()
        pypdf = ModuleType("pypdf")
        pypdf.PdfReader = lambda _: SimpleNamespace(
            pages=[
                SimpleNamespace(
                    images=[
                        SimpleNamespace(
                            data=b"x" * (DocumentService._MAX_IMAGE_BYTES + 1),
                            name="large.png",
                        )
                    ]
                )
            ]
        )

        with patch.dict(
            sys.modules, {"pdfplumber": pdfplumber, "pypdf": pypdf}
        ):
            content = DocumentService.parse(
                UploadedRequirement("需求.pdf", b"fake-pdf")
            )

        self.assertEqual(content.attachments, ())
        self.assertEqual(content.stats.image_count, 0)
        self.assertEqual(content.stats.skipped_image_count, 1)
        self.assertIn(
            DocumentWarningCode.IMAGE_TOO_LARGE,
            [warning.code for warning in content.warnings],
        )

    def test_pdf_table_failure_is_reported_without_losing_page_text(self):
        def fail_tables():
            raise RuntimeError("table parser failed")

        class Pdf:
            pages = [
                SimpleNamespace(
                    extract_text=lambda: "订单需求",
                    extract_tables=fail_tables,
                    curves=[],
                )
            ]

            def __enter__(self):
                return self

            def __exit__(self, *_):
                return None

        pdfplumber = ModuleType("pdfplumber")
        pdfplumber.open = lambda _: Pdf()
        pypdf = ModuleType("pypdf")
        pypdf.PdfReader = lambda _: SimpleNamespace(
            pages=[SimpleNamespace(images=[])]
        )

        with patch.dict(
            sys.modules, {"pdfplumber": pdfplumber, "pypdf": pypdf}
        ):
            content = DocumentService.parse(
                UploadedRequirement("需求.pdf", b"fake-pdf")
            )

        self.assertEqual(content.to_plain_text(), "订单需求")
        self.assertEqual(content.stats.skipped_table_count, 1)
        self.assertIn(
            DocumentWarningCode.TABLE_EXTRACTION_FAILED,
            [warning.code for warning in content.warnings],
        )

    def test_scanned_pdf_ocr_keeps_confidence_source_and_low_confidence_candidate(self):
        from PIL import Image

        class FakeOcr:
            def recognize(self, image_bytes, mime_type):
                return (
                    OcrTextLine("退款金额不得超过订单金额", 0.96),
                    OcrTextLine("疑似模糊审批规则", 0.55),
                )

        class Pdf:
            pages = [
                SimpleNamespace(
                    extract_text=lambda: "",
                    extract_tables=lambda: [],
                    curves=[],
                    to_image=lambda resolution: SimpleNamespace(
                        original=Image.new("RGB", (20, 20), "white")
                    ),
                )
            ]

            def __enter__(self):
                return self

            def __exit__(self, *_):
                return None

        pdfplumber = ModuleType("pdfplumber")
        pdfplumber.open = lambda _: Pdf()
        pypdf = ModuleType("pypdf")
        pypdf.PdfReader = lambda _: SimpleNamespace(
            pages=[SimpleNamespace(images=[])]
        )

        with patch.dict(
            sys.modules, {"pdfplumber": pdfplumber, "pypdf": pypdf}
        ):
            content = DocumentService.parse(
                UploadedRequirement("扫描需求.pdf", b"fake-pdf"),
                ocr_engine=FakeOcr(),
            )

        ocr_elements = [
            element
            for element in content.elements
            if isinstance(element, DocumentOcrElement)
        ]
        self.assertEqual(len(ocr_elements), 2)
        self.assertEqual(ocr_elements[0].source.page_number, 1)
        self.assertEqual(ocr_elements[0].confidence, 0.96)
        self.assertEqual(
            ocr_elements[0].disposition, DocumentOcrDisposition.ACCEPTED
        )
        self.assertEqual(
            ocr_elements[1].disposition,
            DocumentOcrDisposition.REVIEW_REQUIRED,
        )
        self.assertEqual(ocr_elements[0].image_id, content.elements[0].image.image_id)
        self.assertIn("退款金额不得超过订单金额", content.to_plain_text())
        self.assertNotIn("疑似模糊审批规则", content.to_plain_text())
        self.assertEqual(content.stats.ocr_element_count, 2)
        self.assertEqual(content.stats.low_confidence_ocr_count, 1)
        self.assertIn(
            DocumentWarningCode.OCR_LOW_CONFIDENCE,
            [warning.code for warning in content.warnings],
        )

    def test_single_image_ocr_failure_does_not_abort_remaining_docx_images(self):
        from docx import Document
        from PIL import Image

        class FlakyOcr:
            calls = 0

            def recognize(self, image_bytes, mime_type):
                self.calls += 1
                if self.calls == 1:
                    raise OcrError("first image failed")
                return (OcrTextLine("第二张图中的有效规则", 0.93),)

        def image_bytes(color):
            buffer = BytesIO()
            Image.new("RGB", (8, 8), color).save(buffer, format="PNG")
            buffer.seek(0)
            return buffer

        document = Document()
        document.add_paragraph().add_run().add_picture(image_bytes("red"))
        document.add_paragraph().add_run().add_picture(image_bytes("blue"))
        payload = BytesIO()
        document.save(payload)

        content = DocumentService.parse(
            UploadedRequirement("图文需求.docx", payload.getvalue()),
            ocr_engine=FlakyOcr(),
        )

        self.assertEqual(content.stats.image_count, 2)
        self.assertEqual(content.stats.ocr_element_count, 1)
        self.assertEqual(content.stats.failed_ocr_count, 1)
        self.assertIn("第二张图中的有效规则", content.to_plain_text())
        self.assertIn(
            DocumentWarningCode.OCR_FAILED,
            [warning.code for warning in content.warnings],
        )

    def test_visual_candidate_is_analyzed_and_keeps_structured_source(self):
        from docx import Document
        from PIL import Image

        class EmptyOcr:
            def recognize(self, image_bytes, mime_type):
                return ()

        class FakeVisual:
            calls = 0

            def analyze(self, image_bytes, mime_type, *, context, ocr_text):
                self.calls += 1
                self.context = context
                return VisualUnderstandingResult(
                    kind=DocumentVisualKind.FLOWCHART,
                    summary="用户提交后进入风控判断",
                    confidence=0.94,
                    nodes=(
                        DocumentVisualNode("submit", "提交", "action"),
                        DocumentVisualNode("risk", "风控判断", "decision"),
                    ),
                    relations=(
                        DocumentVisualRelation(
                            "submit", "risk", condition="请求有效"
                        ),
                    ),
                    state_changes=("待处理变为审核中",),
                )

        image = BytesIO()
        Image.new("RGB", (400, 260), "white").save(image, format="PNG")
        image.seek(0)
        document = Document()
        paragraph = document.add_paragraph("提现业务流程图")
        paragraph.add_run().add_picture(image)
        payload = BytesIO()
        document.save(payload)
        visual = FakeVisual()

        content = DocumentService.parse(
            UploadedRequirement("流程需求.docx", payload.getvalue()),
            ocr_engine=EmptyOcr(),
            visual_engine=visual,
        )

        visual_elements = [
            element
            for element in content.elements
            if isinstance(element, DocumentVisualElement)
        ]
        self.assertEqual(visual.calls, 1)
        self.assertIn("提现业务流程图", visual.context)
        self.assertEqual(len(visual_elements), 1)
        self.assertEqual(
            visual_elements[0].analysis.kind, DocumentVisualKind.FLOWCHART
        )
        self.assertEqual(
            visual_elements[0].analysis.relations[0].condition, "请求有效"
        )
        self.assertEqual(
            visual_elements[0].analysis.image_id,
            next(
                element.image.image_id
                for element in content.elements
                if isinstance(element, DocumentImageElement)
            ),
        )
        self.assertIn("用户提交后进入风控判断", content.to_plain_text())
        self.assertEqual(content.stats.visual_candidate_count, 1)
        self.assertEqual(content.stats.visual_analyzed_count, 1)
        self.assertEqual(content.stats.failed_visual_count, 0)

    def test_decorative_or_text_only_image_does_not_call_visual_engine(self):
        from docx import Document
        from PIL import Image

        class TextOcr:
            def recognize(self, image_bytes, mime_type):
                return (OcrTextLine("退款金额不得超过订单金额", 0.96),)

        class UnexpectedVisual:
            def analyze(self, *args, **kwargs):
                raise AssertionError("decorative image must not use vision")

        image = BytesIO()
        Image.new("RGB", (400, 260), "white").save(image, format="PNG")
        image.seek(0)
        document = Document()
        paragraph = document.add_paragraph("退款规则说明")
        paragraph.add_run().add_picture(image)
        payload = BytesIO()
        document.save(payload)

        content = DocumentService.parse(
            UploadedRequirement("文字截图.docx", payload.getvalue()),
            ocr_engine=TextOcr(),
            visual_engine=UnexpectedVisual(),
        )

        self.assertEqual(content.stats.visual_candidate_count, 0)
        self.assertEqual(content.stats.visual_analyzed_count, 0)
        self.assertIn("退款金额不得超过订单金额", content.to_plain_text())

    def test_visual_call_limit_and_single_image_failure_are_isolated(self):
        from docx import Document
        from PIL import Image

        class EmptyOcr:
            def recognize(self, image_bytes, mime_type):
                return ()

        class FlakyVisual:
            calls = 0

            def analyze(self, image_bytes, mime_type, *, context, ocr_text):
                self.calls += 1
                if self.calls == 1:
                    raise RuntimeError("first visual request failed")
                return VisualUnderstandingResult(
                    kind=DocumentVisualKind.UI_MOCKUP,
                    summary=f"页面原型{self.calls}",
                    confidence=0.9,
                )

        document = Document()
        for index in range(6):
            image = BytesIO()
            Image.new("RGB", (400, 260), "white").save(
                image, format="PNG"
            )
            image.seek(0)
            paragraph = document.add_paragraph(f"第{index + 1}张UI原型")
            paragraph.add_run().add_picture(image)
        payload = BytesIO()
        document.save(payload)
        visual = FlakyVisual()

        content = DocumentService.parse(
            UploadedRequirement("多图需求.docx", payload.getvalue()),
            ocr_engine=EmptyOcr(),
            visual_engine=visual,
        )

        self.assertEqual(visual.calls, 5)
        self.assertEqual(content.stats.visual_candidate_count, 6)
        self.assertEqual(content.stats.visual_analyzed_count, 4)
        self.assertEqual(content.stats.failed_visual_count, 1)
        codes = [warning.code for warning in content.warnings]
        self.assertIn(DocumentWarningCode.VISION_FAILED, codes)
        self.assertIn(DocumentWarningCode.VISION_LIMIT_EXCEEDED, codes)

    def test_low_confidence_visual_result_stays_out_of_requirement_text(self):
        from docx import Document
        from PIL import Image

        class EmptyOcr:
            def recognize(self, image_bytes, mime_type):
                return ()

        class LowConfidenceVisual:
            def analyze(self, image_bytes, mime_type, *, context, ocr_text):
                return VisualUnderstandingResult(
                    kind=DocumentVisualKind.STATE_DIAGRAM,
                    summary="疑似审核状态变化",
                    confidence=0.69,
                )

        image = BytesIO()
        Image.new("RGB", (400, 260), "white").save(image, format="PNG")
        image.seek(0)
        document = Document()
        paragraph = document.add_paragraph("订单状态图")
        paragraph.add_run().add_picture(image)
        payload = BytesIO()
        document.save(payload)

        content = DocumentService.parse(
            UploadedRequirement("状态需求.docx", payload.getvalue()),
            ocr_engine=EmptyOcr(),
            visual_engine=LowConfidenceVisual(),
        )

        self.assertEqual(content.stats.visual_analyzed_count, 1)
        self.assertNotIn("疑似审核状态变化", content.to_plain_text())
        self.assertIn(
            DocumentWarningCode.VISION_LOW_CONFIDENCE,
            [warning.code for warning in content.warnings],
        )

    def test_pdf_vector_flow_candidate_is_rendered_for_visual_analysis(self):
        from PIL import Image

        class EmptyOcr:
            def recognize(self, image_bytes, mime_type):
                return ()

        class FakeVisual:
            calls = 0

            def analyze(self, image_bytes, mime_type, *, context, ocr_text):
                self.calls += 1
                return VisualUnderstandingResult(
                    kind=DocumentVisualKind.FLOWCHART,
                    summary="审批通过后完成",
                    confidence=0.9,
                )

        class Pdf:
            pages = [
                SimpleNamespace(
                    extract_text=lambda: "审批业务流程图",
                    extract_tables=lambda: [],
                    curves=[{"object_type": "curve"}],
                    to_image=lambda resolution: SimpleNamespace(
                        original=Image.new("RGB", (400, 260), "white")
                    ),
                )
            ]

            def __enter__(self):
                return self

            def __exit__(self, *_):
                return None

        pdfplumber = ModuleType("pdfplumber")
        pdfplumber.open = lambda _: Pdf()
        pypdf = ModuleType("pypdf")
        pypdf.PdfReader = lambda _: SimpleNamespace(
            pages=[SimpleNamespace(images=[])]
        )
        visual = FakeVisual()

        with patch.dict(
            sys.modules, {"pdfplumber": pdfplumber, "pypdf": pypdf}
        ):
            content = DocumentService.parse(
                UploadedRequirement("流程需求.pdf", b"fake-pdf"),
                ocr_engine=EmptyOcr(),
                visual_engine=visual,
            )

        self.assertEqual(visual.calls, 1)
        self.assertEqual(content.stats.visual_candidate_count, 1)
        self.assertEqual(content.stats.visual_analyzed_count, 1)
        self.assertIn("审批通过后完成", content.to_plain_text())

    def test_unsupported_format_and_parser_errors_are_explicit(self):
        with self.assertRaisesRegex(ValueError, "不支持的文件格式"):
            DocumentService.parse(UploadedRequirement("需求.xlsx", b"data"))

        with self.assertRaisesRegex(ValueError, "UnicodeDecodeError"):
            DocumentService.parse(UploadedRequirement("需求.txt", b"\xff"))


if __name__ == "__main__":
    unittest.main()
