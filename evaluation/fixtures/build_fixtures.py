"""Build deterministic synthetic document fixtures for parsing evaluation."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ASSET_DIR = Path(__file__).resolve().parent / "assets"
BODY_FONT = "SyntheticCJK"


def _font_path(*, bold: bool = False) -> Path:
    configured = os.getenv("EVAL_CJK_FONT")
    candidates = [
        Path(configured) if configured else None,
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for candidate in candidates:
        if candidate and candidate.exists():
            return candidate
    raise RuntimeError("A CJK font is required; set EVAL_CJK_FONT")


def _pil_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(_font_path(bold=bold)), size=size)


def _set_run_font(run, size: int, *, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(sum(widths)))
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    properties.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            cell.width = Inches(value / 1440)
            cell._tc.tcPr.tcW.set(qn("w:type"), "dxa")
            cell._tc.tcPr.tcW.set(qn("w:w"), str(value))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_native_text_pdf() -> None:
    path = ASSET_DIR / "native_order_requirement.pdf"
    pdfmetrics.registerFont(TTFont(BODY_FONT, str(_font_path())))
    pdf = canvas.Canvas(str(path), pagesize=letter)
    pdf.setTitle("Synthetic Order Requirement")
    pdf.setFont(BODY_FONT, 18)
    pdf.drawString(72, 720, "订单创建需求")
    pdf.setFont(BODY_FONT, 11)
    lines = [
        "用户提交订单时校验可售库存。",
        "库存充足时创建订单并扣减库存。",
        "库存不足时拒绝创建，并提示“库存不足”。",
        "相同 request_id 重复提交时只能创建一张订单。",
    ]
    for index, line in enumerate(lines):
        pdf.drawString(72, 675 - index * 28, line)
    pdf.save()


def build_role_matrix_docx() -> None:
    path = ASSET_DIR / "role_permission_matrix.docx"
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = Inches(1)
    section.bottom_margin = section.left_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading = document.styles["Heading 1"]
    heading.font.name = "Microsoft YaHei"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(10)

    document.add_heading("项目角色权限矩阵", level=1)
    document.add_paragraph("角色变更后，新权限立即生效。")
    rows = [
        ("角色", "查看项目", "修改项目", "管理成员"),
        ("管理员", "允许", "允许", "允许"),
        ("编辑者", "允许", "允许", "拒绝"),
        ("查看者", "允许", "拒绝", "拒绝"),
    ]
    table = document.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(paragraph.add_run(value), 10, bold=row_index == 0)
    _set_table_geometry(table, [1800, 2520, 2520, 2520])
    document.save(path)


def build_scanned_pdf() -> None:
    image = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(image)
    draw.text((110, 130), "退款处理规则", font=_pil_font(46, bold=True), fill="#172033")
    lines = [
        "已支付订单允许多次部分退款。",
        "累计退款金额不得超过订单实付金额。",
        "退款成功后记录退款单号和累计退款金额。",
        "退款处理中不得重复提交同一退款申请。",
    ]
    for index, line in enumerate(lines):
        draw.text((110, 250 + index * 82), line, font=_pil_font(30), fill="#172033")
    image.save(ASSET_DIR / "scanned_refund_requirement.pdf", "PDF", resolution=150)


def _box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str) -> None:
    draw.rounded_rectangle(xy, radius=18, fill="#EFF5FF", outline="#2F6FEB", width=3)
    left, top, right, bottom = xy
    bounds = draw.textbbox((0, 0), text, font=_pil_font(28, bold=True))
    x = left + (right - left - (bounds[2] - bounds[0])) / 2
    y = top + (bottom - top - (bounds[3] - bounds[1])) / 2
    draw.text((x, y), text, font=_pil_font(28, bold=True), fill="#172033")


def _arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#52606D", width=5)
    x, y = end
    draw.polygon([(x, y), (x - 15, y - 10), (x - 15, y + 10)], fill="#52606D")


def build_flow_diagram() -> None:
    image = Image.new("RGB", (1200, 700), "white")
    draw = ImageDraw.Draw(image)
    _box(draw, (70, 280, 300, 390), "提交订单")
    _box(draw, (430, 280, 700, 390), "校验库存")
    _box(draw, (850, 100, 1130, 210), "创建订单")
    _box(draw, (850, 470, 1130, 580), "提示库存不足")
    _arrow(draw, (300, 335), (430, 335))
    _arrow(draw, (700, 335), (850, 155))
    _arrow(draw, (700, 335), (850, 525))
    draw.text((710, 205), "库存充足", font=_pil_font(24), fill="#137333")
    draw.text((710, 430), "库存不足", font=_pil_font(24), fill="#B3261E")
    image.save(ASSET_DIR / "order_inventory_flow.png")


def build_upload_ui() -> None:
    image = Image.new("RGB", (1200, 760), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((90, 60, 1110, 700), radius=24, fill="white", outline="#D0D7DE", width=2)
    draw.text((140, 105), "发票上传", font=_pil_font(40, bold=True), fill="#172033")
    draw.text((140, 175), "支持 PDF、JPG、PNG，单个文件最大 10MB", font=_pil_font(24), fill="#52606D")
    draw.rounded_rectangle((140, 235, 1060, 410), radius=18, fill="#F7FAFF", outline="#8BA9D6", width=3)
    draw.text((420, 280), "拖拽文件到此处", font=_pil_font(28), fill="#52606D")
    draw.rounded_rectangle((465, 340, 735, 395), radius=12, fill="#2F6FEB")
    draw.text((530, 350), "选择文件", font=_pil_font(24, bold=True), fill="white")
    draw.text((145, 450), "文件名", font=_pil_font(22, bold=True), fill="#172033")
    draw.text((650, 450), "大小", font=_pil_font(22, bold=True), fill="#172033")
    draw.text((870, 450), "处理状态", font=_pil_font(22, bold=True), fill="#172033")
    draw.line((140, 490, 1060, 490), fill="#D0D7DE", width=2)
    draw.text((145, 515), "invoice-001.pdf", font=_pil_font(22), fill="#172033")
    draw.text((650, 515), "2.4MB", font=_pil_font(22), fill="#172033")
    draw.text((870, 515), "上传成功", font=_pil_font(22), fill="#137333")
    draw.text((145, 590), "invoice-002.exe", font=_pil_font(22), fill="#172033")
    draw.text((650, 590), "1.2MB", font=_pil_font(22), fill="#172033")
    draw.text((870, 590), "格式不支持", font=_pil_font(22), fill="#B3261E")
    image.save(ASSET_DIR / "invoice_upload_ui.png")


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        _set_run_font(run, 16 if level == 1 else 13, bold=True, color="1F4E79")


def _add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    _set_run_font(paragraph.add_run(text), 10)


def _add_rule_table(document: Document, rows: list[tuple[str, ...]], widths: list[int]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            _set_run_font(paragraph.add_run(value), 9, bold=row_index == 0)
    _set_table_geometry(table, widths)


def _build_lock_flow_image(path: Path) -> None:
    image = Image.new("RGB", (1100, 520), "white")
    draw = ImageDraw.Draw(image)
    boxes = [
        ((40, 205, 230, 305), "靠近门锁"),
        ((320, 205, 530, 305), "校验距离"),
        ((650, 60, 910, 160), "双向认证"),
        ((650, 350, 910, 450), "保持锁定"),
        ((950, 60, 1070, 160), "开锁"),
    ]
    for xy, text in boxes:
        _box(draw, xy, text)
    _arrow(draw, (230, 255), (320, 255))
    _arrow(draw, (530, 255), (650, 110))
    _arrow(draw, (530, 255), (650, 400))
    _arrow(draw, (910, 110), (950, 110))
    draw.text((535, 135), "距离≤1.5米", font=_pil_font(21), fill="#137333")
    draw.text((535, 350), "距离>1.7米", font=_pil_font(21), fill="#B3261E")
    image.save(path)


def build_realistic_smart_lock_docx() -> None:
    flow_path = ASSET_DIR / "smart_lock_flow.png"
    _build_lock_flow_image(flow_path)
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.85)
    _add_heading(document, "智能门锁蓝牙近场自动解锁需求", 1)
    _add_body(document, "版本：V1.3　产品：HomeLink 智能门锁　文档性质：脱敏合成评测样本")
    _add_heading(document, "1. 背景与目标", 2)
    _add_body(document, "用户携带已绑定手机靠近门锁时，系统通过低功耗蓝牙完成距离判断和双向认证，在满足安全条件后自动开锁，减少手动操作。")
    _add_heading(document, "2. 核心业务规则", 2)
    _add_rule_table(document, [
        ("规则", "触发条件", "预期结果"),
        ("近场候选", "测距结果不大于1.5米", "进入双向认证"),
        ("远离保护", "测距结果大于1.7米", "保持锁定，不发起认证"),
        ("认证成功", "手机已绑定且动态凭证有效", "驱动门锁并记录成功日志"),
        ("连续失败", "5分钟内认证失败达到3次", "暂停自动解锁10分钟"),
    ], [1800, 3300, 4260])
    _add_heading(document, "3. 状态与异常", 2)
    _add_body(document, "手机蓝牙关闭、门锁离线或电量低于10%时不得自动开锁，并向App展示明确原因。暂停期间即使距离满足也不得发起自动解锁。")
    _add_body(document, "同一动态凭证只能成功使用一次；重复请求不得造成重复开锁或重复成功日志。")
    _add_heading(document, "4. 自动解锁流程图", 2)
    document.add_picture(str(flow_path), width=Inches(6.3))
    caption = document.add_paragraph("图1　近场判断与认证主流程")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_heading(document, "5. 待确认事项", 2)
    _add_body(document, "1）1.5米至1.7米灰区内如何处理；2）暂停期结束后是否自动恢复；3）低电量阈值是否允许由管理员配置。")
    document.save(ASSET_DIR / "realistic_smart_lock_prd.docx")


def _build_coupon_ui_image(path: Path) -> None:
    image = Image.new("RGB", (1100, 620), "#F6F8FB")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((70, 45, 1030, 575), radius=24, fill="white", outline="#C9D4E5", width=3)
    draw.text((110, 80), "订单结算", font=_pil_font(36, bold=True), fill="#172033")
    draw.text((110, 155), "商品金额", font=_pil_font(24), fill="#52606D")
    draw.text((830, 155), "¥239.00", font=_pil_font(24, bold=True), fill="#172033")
    draw.text((110, 220), "优惠券", font=_pil_font(24), fill="#52606D")
    draw.rounded_rectangle((420, 205, 940, 265), radius=10, fill="#FFF7E6", outline="#F3B23C")
    draw.text((445, 220), "满200减30（剩余14分钟）", font=_pil_font(22), fill="#8A5300")
    draw.text((110, 315), "应付金额", font=_pil_font(28, bold=True), fill="#172033")
    draw.text((800, 310), "¥209.00", font=_pil_font(30, bold=True), fill="#D93025")
    draw.rounded_rectangle((690, 420, 950, 500), radius=14, fill="#2F6FEB")
    draw.text((760, 440), "提交订单", font=_pil_font(26, bold=True), fill="white")
    image.save(path)


def build_realistic_checkout_docx() -> None:
    ui_path = ASSET_DIR / "checkout_coupon_ui.png"
    _build_coupon_ui_image(ui_path)
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.85)
    _add_heading(document, "电商订单履约与优惠结算需求 PRD", 1)
    _add_body(document, "版本：V2.2　业务域：交易履约　文档性质：脱敏合成评测样本")
    _add_heading(document, "1. 需求范围", 2)
    _add_body(document, "本期覆盖普通实物商品下单、库存预占、优惠券核销、支付结果处理和订单取消，不覆盖预售、虚拟商品、跨境订单和多张优惠券叠加。")
    _add_heading(document, "2. 优惠与金额规则", 2)
    _add_rule_table(document, [
        ("规则项", "规则说明", "失败提示"),
        ("满减券", "商品金额满200元减30元，运费不计入门槛", "未达到使用门槛"),
        ("使用数量", "每笔订单最多使用一张优惠券", "每单限用一张"),
        ("有效期", "提交订单时校验有效期，支付前再次校验", "优惠券已失效"),
        ("金额精度", "人民币金额保留两位小数", "-"),
    ], [1600, 5100, 2660])
    _add_heading(document, "3. 库存与幂等", 2)
    _add_body(document, "提交订单时按SKU校验可售库存。全部SKU充足才创建订单并预占库存；任一SKU不足则整单失败且不得核销优惠券。相同request_id重复或并发提交只能创建一张有效订单，库存和优惠券均只处理一次。")
    _add_heading(document, "4. 支付与取消", 2)
    _add_body(document, "支付成功回调需记录支付平台流水号。重复回调不得重复记账。待支付订单可由用户取消，取消成功后释放库存并触发优惠券退回；优惠券退回后的可用状态取决于其是否已经过期。")
    _add_heading(document, "5. 结算页交互原型", 2)
    document.add_picture(str(ui_path), width=Inches(6.2))
    caption = document.add_paragraph("图1　优惠券选中后的结算页示意")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_heading(document, "6. 待确认事项", 2)
    _add_body(document, "1）支付过程中优惠券过期时是继续支付还是阻断；2）取消后已过期优惠券的最终状态；3）库存预占有效期和超时释放时机。")
    document.save(ASSET_DIR / "realistic_checkout_prd.docx")


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    build_native_text_pdf()
    build_role_matrix_docx()
    build_scanned_pdf()
    build_flow_diagram()
    build_upload_ui()
    build_realistic_smart_lock_docx()
    build_realistic_checkout_docx()
    print(f"generated 9 fixture assets in {ASSET_DIR}")


if __name__ == "__main__":
    main()
